#!/usr/bin/env python3
"""
Candidate Screening Engine (`screener_core.py`)
Evaluates candidate CVs against Job Description using Google GenAI SDK (gemini-2.5-flash),
applies HR Natural Break / Gap Analysis Tiering, and exports Markdown, Excel, and Word summaries.
"""

import sys
import os
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

# Unbuffer stdout and reconfigure for Windows UTF-8 console output
os.environ["PYTHONUNBUFFERED"] = "1"
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr and hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# Third-party dependencies
import pypdf
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import docx.opc.constants
import docx.shared
import docx.enum.text

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from google import genai
from google.genai import types
from pydantic import BaseModel, Field


# ---------------------------------------------------------------------------
# API Key Resolution Helper (Self-Healing)
# ---------------------------------------------------------------------------
def resolve_api_key(target_dir: Optional[Path] = None) -> Optional[str]:
    """
    Search for GEMINI_API_KEY or GOOGLE_API_KEY in:
    1. OS environment variables
    2. .env files in script directory, target directory (and parents), or current working directory (and parents)
    3. Interactive prompt fallback with self-healing persistence to .env at script root
    """
    for env_var in ["GEMINI_API_KEY", "GOOGLE_API_KEY"]:
        if os.environ.get(env_var):
            return os.environ[env_var]

    if getattr(sys, 'frozen', False):
        script_dir = Path(sys.executable).resolve().parent
    else:
        script_dir = Path(__file__).resolve().parent
    search_paths = [script_dir]
    search_paths.extend(script_dir.parents)
    try:
        if script_dir.parent.exists():
            for sib in script_dir.parent.iterdir():
                if sib.is_dir():
                    search_paths.append(sib)
    except Exception:
        pass

    if target_dir:
        search_paths.append(target_dir)
        search_paths.extend(target_dir.parents)
    
    curr = Path.cwd().resolve()
    search_paths.append(curr)
    search_paths.extend(curr.parents)

    visited = set()
    for p in search_paths:
        try:
            p_resolved = p.resolve()
        except Exception:
            continue
        if p_resolved in visited:
            continue
        visited.add(p_resolved)
        for env_name in [".env", ".env.local"]:
            env_file = p_resolved / env_name
            if env_file.exists():
                try:
                    content = env_file.read_text(encoding="utf-8", errors="ignore")
                    for line in content.splitlines():
                        line = line.strip()
                        if "=" in line and not line.startswith("#"):
                            k, v = line.split("=", 1)
                            k = k.strip()
                            v = v.strip().strip('"').strip("'")
                            if k in ["GEMINI_API_KEY", "GOOGLE_API_KEY"] and v:
                                os.environ["GEMINI_API_KEY"] = v
                                return v
                except Exception:
                    pass

    # API Key Self-Healing Prompt
    print("\n============================================================")
    print("🔑 GEMINI_API_KEY / GOOGLE_API_KEY not found in environment or .env file.")
    print("============================================================")
    try:
        user_key = input("Please enter your GEMINI_API_KEY: ").strip()
    except (KeyboardInterrupt, EOFError):
        user_key = ""

    if user_key:
        env_dest = script_dir / ".env"
        try:
            env_dest.write_text(f"GEMINI_API_KEY={user_key}\n", encoding="utf-8")
            print(f"✔ API Key saved to '{env_dest.name}' at script root.")
        except Exception as e:
            print(f"  [Warning] Could not persist API key to .env: {e}")
        os.environ["GEMINI_API_KEY"] = user_key
        return user_key

    return None


# ---------------------------------------------------------------------------
# Pydantic Schema for Structured GenAI Response
# ---------------------------------------------------------------------------
class CandidateEvaluation(BaseModel):
    candidate_name: str = Field(description="Full name of candidate extracted from CV")
    score: int = Field(description="Overall fit score from 0 to 100 based on Job Description requirements")
    experience_summary: str = Field(description="Summary of work experience and total estimated years of experience")
    key_strengths: List[str] = Field(description="Top key strengths directly matching JD requirements")
    key_gaps: List[str] = Field(description="Key missing qualifications, skill gaps, or weaknesses relative to JD")
    education_certs: str = Field(description="Degrees and relevant professional certifications")
    executive_summary: str = Field(description="Concise 2-3 sentence executive recommendation for HR decision")


# ---------------------------------------------------------------------------
# File Text Extraction Utilities
# ---------------------------------------------------------------------------
def extract_text_from_pdf(filepath: Path) -> tuple[str, bool, Optional[bytes]]:
    """
    Extract text from PDF file using pypdf.
    Returns tuple: (extracted_text, is_scanned_pdf, raw_bytes)
    Handles zero-byte PDFs, encrypted PDFs, and scanned image PDFs.
    """
    if not filepath.exists():
        print(f"  [Warning] File '{filepath.name}' does not exist.")
        return "", False, None

    file_size = filepath.stat().st_size
    if file_size == 0:
        print(f"  [Warning] PDF file '{filepath.name}' is 0 bytes (zero-byte file).")
        return "[Warning: Zero-byte PDF file - empty content]", False, None

    raw_bytes = None
    try:
        raw_bytes = filepath.read_bytes()
    except Exception as e:
        print(f"  [Warning] Could not read bytes of PDF '{filepath.name}': {e}")

    text_chunks = []
    try:
        reader = pypdf.PdfReader(str(filepath))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                pass

        for page_idx, page in enumerate(reader.pages):
            try:
                t = page.extract_text()
                if t:
                    text_chunks.append(t.strip())
            except Exception as pe:
                print(f"  [Warning] Page {page_idx+1} text extraction failed in '{filepath.name}': {pe}")
    except Exception as e:
        print(f"  [Warning] Error parsing PDF structure '{filepath.name}': {e}")

    full_text = "\n".join(text_chunks).strip()

    # Detect Scanned / Image-only PDF (file size > 2KB but extracted text is minimal/empty)
    is_scanned = False
    if file_size > 2048 and len(full_text) < 30:
        is_scanned = True
        print(f"  [Notice] Scanned/Image PDF detected: '{filepath.name}' ({file_size} bytes). Multimodal GenAI OCR active.")
        if not full_text:
            full_text = f"[Scanned/Image PDF Document: '{filepath.name}' - Native Multimodal GenAI OCR active]"

    return full_text, is_scanned, raw_bytes


def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from Word DOCX file using python-docx."""
    if not filepath.exists() or filepath.stat().st_size == 0:
        print(f"  [Warning] DOCX file '{filepath.name}' is missing or 0 bytes.")
        return ""

    text_chunks = []
    try:
        doc = docx.Document(str(filepath))
        for p in doc.paragraphs:
            if p.text.strip():
                text_chunks.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))
    except Exception as e:
        print(f"  [Warning] Error reading DOCX '{filepath.name}': {e}")
    return "\n".join(text_chunks)


def extract_text_from_txt(filepath: Path) -> str:
    """Extract text from plain text file with multi-encoding fallback."""
    if not filepath.exists() or filepath.stat().st_size == 0:
        print(f"  [Warning] TXT file '{filepath.name}' is missing or 0 bytes.")
        return ""

    for encoding in ["utf-8", "utf-8-sig", "cp1252", "latin-1", "tis-620", "utf-16"]:
        try:
            return filepath.read_text(encoding=encoding)
        except Exception:
            continue
    return ""


def extract_file_info(filepath: Path) -> Dict[str, Any]:
    """
    Extract text and file metadata from supported document formats (.pdf, .docx, .doc, .txt, .md).
    Returns dict: {"text": str, "is_scanned": bool, "raw_bytes": bytes, "extension": str}
    """
    ext = filepath.suffix.lower()
    if ext == ".pdf":
        text, is_scanned, raw_bytes = extract_text_from_pdf(filepath)
        return {
            "text": text,
            "is_scanned": is_scanned,
            "raw_bytes": raw_bytes,
            "extension": ext
        }
    elif ext in [".docx", ".doc"]:
        text = extract_text_from_docx(filepath)
        return {
            "text": text,
            "is_scanned": False,
            "raw_bytes": None,
            "extension": ext
        }
    elif ext in [".txt", ".md"]:
        text = extract_text_from_txt(filepath)
        return {
            "text": text,
            "is_scanned": False,
            "raw_bytes": None,
            "extension": ext
        }
    return {
        "text": "",
        "is_scanned": False,
        "raw_bytes": None,
        "extension": ext
    }


def extract_file_text(filepath: Path) -> str:
    """Extract raw text from supported document formats."""
    return extract_file_info(filepath)["text"]


def is_placeholder_file(filepath: Path) -> bool:
    """Check if file is a placeholder file based on filename or zero/tiny size."""
    if filepath.stat().st_size == 0:
        return True
    name_lower = filepath.name.lower()
    if any(p in name_lower for p in ["placeholder", "วางไฟล์", "โยนไฟล์", "00_place"]):
        return True
    if filepath.suffix.lower() == ".txt" and filepath.stat().st_size < 50:
        return True
    return False


# ---------------------------------------------------------------------------
# Job Description & Candidate CV Loaders
# ---------------------------------------------------------------------------
def find_and_read_job_description(target_dir: Path) -> tuple[Optional[Path], Dict[str, Any]]:
    """
    Search for Job Description in target folder:
    1. Check `jd/` subfolder first (scanning direct files and subfolders inside `jd/`).
    2. Fallback to target_dir root if `jd/` subfolder is not present or has no valid files.
    """
    excluded_names = {
        "candidate_summary.md", "candidate_summary.xlsx", "candidate_summary.docx",
        "run_screening.bat", "screener_core.py", "requirements.txt"
    }

    supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md"}
    
    # 1. Search in jd/ subfolder
    jd_subfolder = target_dir / "jd"
    candidate_files = []
    
    if jd_subfolder.exists() and jd_subfolder.is_dir():
        for root, dirs, files in os.walk(jd_subfolder):
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            for file_name in files:
                file_path = Path(root) / file_name
                if file_path.suffix.lower() in supported_extensions and not is_placeholder_file(file_path):
                    candidate_files.append(file_path)
        
    # 2. Fallback to target_dir root if jd/ has no valid files
    if not candidate_files:
        candidate_files = [
            f for f in target_dir.iterdir()
            if f.is_file() and f.suffix.lower() in supported_extensions and f.name.lower() not in excluded_names and not is_placeholder_file(f)
        ]

    if not candidate_files:
        return None, {"text": "", "is_scanned": False, "raw_bytes": None, "extension": ""}

    # Prioritize files with jd / job / description in name
    jd_file = None
    for f in candidate_files:
        name_lower = f.name.lower()
        if any(kw in name_lower for kw in ["jd", "job", "description", "ticket", "spec"]):
            jd_file = f
            break
            
    if not jd_file:
        jd_file = candidate_files[0]

    jd_info = extract_file_info(jd_file)
    return jd_file, jd_info


def find_and_read_candidate_cvs(target_dir: Path) -> List[Dict[str, Any]]:
    """
    Recursively scan subdirectories for candidate CV files (.pdf, .docx, .txt, .md).
    Checks `cv/` subfolder first (scanning direct files and subfolders inside `cv/`),
    or fallbacks to any subdirectories of target_dir.
    Excludes placeholder files.
    """
    supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md"}
    cv_list = []

    # Check if cv/ subfolder exists
    cv_subfolder = target_dir / "cv"
    search_root = cv_subfolder if (cv_subfolder.exists() and cv_subfolder.is_dir()) else target_dir

    for root, dirs, files in os.walk(search_root):
        dirs[:] = [d for d in dirs if not d.startswith(".")]
        current_path = Path(root)
        
        if current_path == target_dir:
            # Skip target_dir root files
            continue

        for file_name in files:
            file_path = current_path / file_name
            if file_path.suffix.lower() in supported_extensions and not is_placeholder_file(file_path):
                rel_path = file_path.relative_to(target_dir)
                file_info = extract_file_info(file_path)
                cv_text = file_info["text"]
                if cv_text.strip() or file_info["is_scanned"] or file_info["raw_bytes"]:
                    cv_list.append({
                        "abs_path": file_path,
                        "rel_path": str(rel_path).replace("\\", "/"),
                        "filename": file_name,
                        "text": cv_text,
                        "is_scanned": file_info["is_scanned"],
                        "raw_bytes": file_info["raw_bytes"],
                        "extension": file_info["extension"]
                    })
                else:
                    print(f"  [Warning] Skipping empty CV: {rel_path}")

    return cv_list


# ---------------------------------------------------------------------------
# Gemini AI Evaluation Engine
# ---------------------------------------------------------------------------
def evaluate_cv_against_jd(client: genai.Client, jd_text: str, cv_item: Dict[str, Any], jd_info: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Invoke gemini-2.5-flash with structured JSON response schema, supporting multimodal PDF fallback."""
    prompt = f"""
You are an expert Senior HR Recruiter and Talent Acquisition Specialist.
Evaluate the candidate CV against the target Job Description (JD).

=== TARGET JOB DESCRIPTION ===
{jd_text}

=== CANDIDATE RESUME / CV ({cv_item['rel_path']}) ===
{cv_item['text']}

=== EVALUATION RULES ===
1. Analyze candidate technical skills, domain experience, soft skills, and qualifications vs JD.
2. Determine full name from CV text. If name is not stated, infer from filename '{cv_item['filename']}'.
3. Assign a quantitative fit score from 0 to 100:
   - 90-100: Exceptional fit, exceeds major criteria.
   - 75-89: Strongly qualified, meets core requirements.
   - 50-74: Moderate fit, meets basic needs but has noticeable gaps.
   - 0-49: Unqualified / Poor fit, missing mandatory qualifications.
4. Extract list of key strengths matching JD.
5. Extract list of key gaps / weaknesses relative to JD.
6. Summarize education & certifications.
7. Provide a 2-3 sentence executive recommendation summary.
"""

    contents: List[Any] = [prompt]

    # Multimodal Fallback: Attach raw PDF bytes if candidate CV is scanned / image PDF
    if cv_item.get("raw_bytes") and (cv_item.get("is_scanned") or cv_item.get("extension") == ".pdf"):
        if cv_item.get("is_scanned") or not cv_item.get("text", "").strip() or len(cv_item.get("text", "")) < 100:
            contents.append(
                types.Part.from_bytes(data=cv_item["raw_bytes"], mime_type="application/pdf")
            )

    # Attach raw PDF bytes if JD is scanned / image PDF
    if jd_info and jd_info.get("raw_bytes") and (jd_info.get("is_scanned") or jd_info.get("extension") == ".pdf"):
        if jd_info.get("is_scanned") or not jd_text.strip() or len(jd_text) < 100:
            contents.append(
                types.Part.from_bytes(data=jd_info["raw_bytes"], mime_type="application/pdf")
            )

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=contents,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=CandidateEvaluation,
            temperature=0.2
        )
    )

    result_data = json.loads(response.text)
    result_data["rel_path"] = cv_item["rel_path"]
    result_data["filename"] = cv_item["filename"]
    return result_data


# ---------------------------------------------------------------------------
# HR Natural Break / Gap Analysis Tiering
# ---------------------------------------------------------------------------
def apply_hr_natural_break_tiering(evaluations: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Apply HR Natural Break / Gap Analysis Tiering:
    - Sort candidates by score descending.
    - Floor threshold: score < 50 -> "Rejected"
    - Shortlist: score >= 75 AND gap from previous candidate <= 12 (up to Top 5 max)
    - Longlist: qualified candidates (>= 50) not in Shortlist.
    """
    sorted_evals = sorted(evaluations, key=lambda x: x["score"], reverse=True)
    shortlist_count = 0

    for i, cand in enumerate(sorted_evals):
        score = cand["score"]
        if score < 50:
            cand["tier"] = "Rejected"
        else:
            if score >= 75 and shortlist_count < 5:
                if shortlist_count == 0:
                    cand["tier"] = "Shortlist"
                    shortlist_count += 1
                else:
                    prev_cand = sorted_evals[i - 1]
                    gap = prev_cand["score"] - score
                    if prev_cand["tier"] == "Shortlist" and gap <= 12:
                        cand["tier"] = "Shortlist"
                        shortlist_count += 1
                    else:
                        cand["tier"] = "Longlist"
            else:
                cand["tier"] = "Longlist"

    return sorted_evals


# ---------------------------------------------------------------------------
# Markdown Report Exporter
# ---------------------------------------------------------------------------
def export_candidate_summary_md(target_dir: Path, jd_filename: str, candidates: List[Dict[str, Any]]) -> Path:
    """Generate Candidate_Summary.md executive summary report in 1-Page A4 compact format."""
    md_path = target_dir / "Candidate_Summary.md"
    
    total = len(candidates)
    shortlisted = [c for c in candidates if c["tier"] == "Shortlist"]
    longlisted = [c for c in candidates if c["tier"] == "Longlist"]
    rejected = [c for c in candidates if c["tier"] == "Rejected"]
    
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    md_lines = [
        f"# Candidate Evaluation Executive Summary",
        f"",
        f"**Job Target:** `{target_dir.name}` | **JD:** `{jd_filename}` | **Date:** `{now_str}` | **Total Assessed:** `{total}` (🟢 Shortlist: `{len(shortlisted)}`, 🟡 Longlist: `{len(longlisted)}`, 🔴 Rejected: `{len(rejected)}`)",
        f"",
        f"---",
        f""
    ]
    
    # 1. Shortlist Candidates (Top Priority)
    if shortlisted:
        md_lines.extend([
            f"## 🟢 Shortlist (Top Candidates)",
            f"",
            f"| Rank | Candidate Name | Score | Key Strengths | Key Gaps | Executive Recommendation |",
            f"| :---: | :--- | :---: | :--- | :--- | :--- |"
        ])
        for cand in shortlisted:
            idx = candidates.index(cand) + 1
            cand_name_link = f"[{cand['candidate_name']}]({cand['rel_path']})"
            strengths_str = "<br>• ".join(cand.get("key_strengths", []))
            if strengths_str:
                strengths_str = "• " + strengths_str
            gaps_str = "<br>• ".join(cand.get("key_gaps", []))
            if gaps_str:
                gaps_str = "• " + gaps_str
            exec_summary = cand.get("executive_summary", "").replace("\n", " ")
            md_lines.append(
                f"| {idx} | {cand_name_link} | **{cand['score']}** | {strengths_str} | {gaps_str} | {exec_summary} |"
            )
        md_lines.append("")

    # 2. Longlist Candidates (Backup Pool)
    if longlisted:
        md_lines.extend([
            f"## 🟡 Longlist (Backup Pool)",
            f"",
            f"| Rank | Candidate Name | Score | Brief Summary / Note |",
            f"| :---: | :--- | :---: | :--- |"
        ])
        for cand in longlisted:
            idx = candidates.index(cand) + 1
            cand_name_link = f"[{cand['candidate_name']}]({cand['rel_path']})"
            exec_summary = cand.get("executive_summary", "").replace("\n", " ")
            md_lines.append(
                f"| {idx} | {cand_name_link} | **{cand['score']}** | {exec_summary} |"
            )
        md_lines.append("")

    # 3. Rejected Candidates (Filtered Out < 50)
    if rejected:
        md_lines.extend([
            f"## 🔴 Filtered Out (Score < 50)",
            f""
        ])
        inline_items = []
        for cand in rejected:
            cand_link = f"[{cand['candidate_name']}]({cand['rel_path']})"
            inline_items.append(f"{cand_link} ({cand['score']})")
            
        inline_str = " | ".join(inline_items)
        md_lines.append(f"• {inline_str} — *(Unmatched core qualifications)*")
        md_lines.append("")

    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    return md_path



# ---------------------------------------------------------------------------
# Excel Report Exporter (Harrow Brand Identity)
# ---------------------------------------------------------------------------
def export_candidate_summary_excel(target_dir: Path, jd_filename: str, candidates: List[Dict[str, Any]]) -> Path:
    """Generate Candidate_Summary.xlsx stylized Excel workbook matching Harrow brand identity."""
    xlsx_path = target_dir / "Candidate_Summary.xlsx"
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Screening Summary"
    ws.views.sheetView[0].showGridLines = True
    
    # Harrow Brand Colors & Typography
    # Deep Navy: #00205B | Warm Gold: #C49A45
    header_fill = PatternFill(start_color="00205B", end_color="00205B", fill_type="solid")
    header_font = Font(name="Arial", size=9, bold=True, color="FFFFFF")
    
    title_font = Font(name="Georgia", size=16, bold=True, color="00205B")
    subtitle_font = Font(name="Arial", size=9, italic=True, color="595959")
    footer_font = Font(name="Georgia", size=9, italic=True, color="00205B")
    
    tier_styles = {
        "Shortlist": {
            "fill": PatternFill(start_color="EBF1E5", end_color="EBF1E5", fill_type="solid"),
            "font": Font(name="Arial", size=9, bold=True, color="1E4D2B")
        },
        "Longlist": {
            "fill": PatternFill(start_color="FEF9E7", end_color="FEF9E7", fill_type="solid"),
            "font": Font(name="Arial", size=9, bold=True, color="8D6E18")
        },
        "Rejected": {
            "fill": PatternFill(start_color="FDEDEC", end_color="FDEDEC", fill_type="solid"),
            "font": Font(name="Arial", size=9, bold=True, color="922B21")
        }
    }
    
    link_font = Font(name="Arial", size=9, underline="single", color="0563C1")
    regular_font = Font(name="Arial", size=9)
    bold_font = Font(name="Arial", size=9, bold=True)
    
    gold_border = Border(
        left=Side(style='thin', color='C49A45'),
        right=Side(style='thin', color='C49A45'),
        top=Side(style='thin', color='C49A45'),
        bottom=Side(style='thin', color='C49A45')
    )
    
    # Title Block
    ws.merge_cells("A1:H1")
    ws["A1"] = f"Candidate Evaluation Executive Summary — {target_dir.name}"
    ws["A1"].font = title_font
    
    ws.merge_cells("A2:H2")
    ws["A2"] = f"Job Description: {jd_filename} | Date: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Total Candidates: {len(candidates)}"
    ws["A2"].font = subtitle_font
    
    # Headers
    headers = [
        "Rank", "Candidate Name", "Score", "Tier", 
        "Key Strengths", "Key Gaps", "Experience & Education", "Executive Summary"
    ]
    
    start_row = 4
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = gold_border
    ws.row_dimensions[start_row].height = 28
    
    # Data Rows
    for idx, cand in enumerate(candidates, 1):
        row = start_row + idx
        
        strengths_bullet = "\n• ".join(cand.get("key_strengths", []))
        if strengths_bullet:
            strengths_bullet = "• " + strengths_bullet
            
        gaps_bullet = "\n• ".join(cand.get("key_gaps", []))
        if gaps_bullet:
            gaps_bullet = "• " + gaps_bullet
            
        exp_edu = f"Exp: {cand.get('experience_summary', '')}\nEdu: {cand.get('education_certs', '')}"
        
        # 1. Rank
        rank_cell = ws.cell(row=row, column=1, value=idx)
        rank_cell.alignment = Alignment(horizontal="center", vertical="top")
        rank_cell.font = regular_font
        
        # 2. Candidate Name (Clickable Hyperlink formula for CV)
        name_cell = ws.cell(row=row, column=2)
        name_cell.value = f'=HYPERLINK("{cand["rel_path"]}", "{cand["candidate_name"]}")'
        name_cell.font = link_font
        name_cell.alignment = Alignment(horizontal="left", vertical="top")
        
        # 3. Score
        score_cell = ws.cell(row=row, column=3, value=cand["score"])
        score_cell.font = bold_font
        score_cell.alignment = Alignment(horizontal="center", vertical="top")
        
        # 4. Tier
        tier_cell = ws.cell(row=row, column=4, value=cand["tier"])
        tier_cell.alignment = Alignment(horizontal="center", vertical="top")
        t_style = tier_styles.get(cand["tier"], {})
        if "fill" in t_style:
            tier_cell.fill = t_style["fill"]
        if "font" in t_style:
            tier_cell.font = t_style["font"]
            
        # 5. Strengths
        str_cell = ws.cell(row=row, column=5, value=strengths_bullet)
        str_cell.font = regular_font
        str_cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        
        # 6. Gaps
        gap_cell = ws.cell(row=row, column=6, value=gaps_bullet)
        gap_cell.font = regular_font
        gap_cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        
        # 7. Experience & Education
        exp_cell = ws.cell(row=row, column=7, value=exp_edu)
        exp_cell.font = regular_font
        exp_cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        
        # 8. Executive Summary
        exec_cell = ws.cell(row=row, column=8, value=cand.get("executive_summary", ""))
        exec_cell.font = regular_font
        exec_cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        
        for c in range(1, 9):
            ws.cell(row=row, column=c).border = gold_border
                
        ws.row_dimensions[row].height = 65
        
    # Safeguarding Footer Notice
    footer_row = start_row + len(candidates) + 2
    ws.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=8)
    footer_cell = ws.cell(row=footer_row, column=1, value="Harrow International School Bangkok is committed to the safety and protection of children.")
    footer_cell.font = footer_font
    footer_cell.alignment = Alignment(horizontal="center", vertical="center")
    
    col_widths = {
        1: 8,   # Rank
        2: 26,  # Candidate Name
        3: 10,  # Score
        4: 14,  # Tier
        5: 38,  # Strengths
        6: 38,  # Gaps
        7: 32,  # Experience & Education
        8: 45   # Executive Summary
    }
    
    for col_idx, width in col_widths.items():
        col_letter = get_column_letter(col_idx)
        ws.column_dimensions[col_letter].width = width
        
    wb.save(xlsx_path)
    return xlsx_path


# ---------------------------------------------------------------------------
# DOCX Report Exporter (Harrow Brand Identity)
# ---------------------------------------------------------------------------
def set_cell_background(cell, fill_hex: str):
    """Set cell background shading in python-docx."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_hyperlink(paragraph, url: str, text: str, color="0563C1", underline=True):
    """Add a clickable relative hyperlink to a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w", "r")} r:id="{r_id}"/>')
    u_val = "single" if underline else "none"
    new_run = parse_xml(f'<w:r {nsdecls("w", "r")}><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="{color}"/><w:u w:val="{u_val}"/></w:rPr><w:t>{text}</w:t></w:r>')
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def export_candidate_summary_docx(target_dir: Path, jd_filename: str, candidates: List[Dict[str, Any]]) -> Path:
    """Generate Candidate_Summary.docx executive summary Word document matching 1-page A4 Candidate_Summary.md layout."""
    docx_path = target_dir / "Candidate_Summary.docx"
    doc = docx.Document()
    
    # Page Margins (Narrow margins to fit 1-page A4)
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(0.5)
        section.bottom_margin = docx.shared.Inches(0.5)
        section.left_margin = docx.shared.Inches(0.5)
        section.right_margin = docx.shared.Inches(0.5)
        
        # Document Footer Safeguarding Notice
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("Harrow International School Bangkok is committed to the safety and protection of children.")
        footer_run.font.name = "Georgia"
        footer_run.font.size = docx.shared.Pt(8.5)
        footer_run.font.italic = True
        footer_run.font.color.rgb = docx.shared.RGBColor(0, 32, 91)

    total = len(candidates)
    shortlisted = [c for c in candidates if c["tier"] == "Shortlist"]
    longlisted = [c for c in candidates if c["tier"] == "Longlist"]
    rejected = [c for c in candidates if c["tier"] == "Rejected"]
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Header 1: Document Title (Georgia 16pt Bold Deep Navy #00205B)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = docx.shared.Pt(2)
    title_run = title_p.add_run("Candidate Evaluation Executive Summary")
    title_run.font.name = "Georgia"
    title_run.font.size = docx.shared.Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = docx.shared.RGBColor(0, 32, 91)
    
    # Subheader / Meta Info (Single italicized line)
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = docx.shared.Pt(8)
    meta_str = f"Job Target: {target_dir.name}  |  JD: {jd_filename}  |  Date: {now_str}  |  Total Assessed: {total} (🟢 Shortlist: {len(shortlisted)}, 🟡 Longlist: {len(longlisted)}, 🔴 Rejected: {len(rejected)})"
    meta_run = sub_p.add_run(meta_str)
    meta_run.font.name = "Arial"
    meta_run.font.size = docx.shared.Pt(8.5)
    meta_run.font.italic = True
    meta_run.font.color.rgb = docx.shared.RGBColor(89, 89, 89)

    # 1. Shortlist Section (Top Priority)
    if shortlisted:
        h2_1 = doc.add_paragraph()
        h2_1.paragraph_format.space_before = docx.shared.Pt(6)
        h2_1.paragraph_format.space_after = docx.shared.Pt(4)
        run_1 = h2_1.add_run("🟢 Shortlist (Top Candidates)")
        run_1.font.name = "Georgia"
        run_1.font.size = docx.shared.Pt(11)
        run_1.font.bold = True
        run_1.font.color.rgb = docx.shared.RGBColor(0, 32, 91)
        
        t_short = doc.add_table(rows=len(shortlisted)+1, cols=6)
        t_short.style = 'Table Grid'
        
        headers_short = ["Rank", "Candidate Name", "Score", "Key Strengths", "Key Gaps", "Executive Recommendation"]
        hdr_cells = t_short.rows[0].cells
        for i, h in enumerate(headers_short):
            hdr_cells[i].text = h
            set_cell_background(hdr_cells[i], "00205B")
            p = hdr_cells[i].paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = docx.shared.Pt(8.5)
                r.font.bold = True
                r.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
                
        for s_idx, cand in enumerate(shortlisted, 1):
            row_idx = s_idx
            r_cells = t_short.rows[row_idx].cells
            
            # Rank
            r_cells[0].text = str(candidates.index(cand) + 1)
            r_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Name
            p_name = r_cells[1].paragraphs[0]
            add_hyperlink(p_name, cand["rel_path"], cand["candidate_name"])
            
            # Score
            r_cells[2].text = str(cand["score"])
            r_cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            r_cells[2].paragraphs[0].runs[0].font.bold = True
            
            # Strengths
            r_cells[3].text = "• " + "\n• ".join(cand.get("key_strengths", [])) if cand.get("key_strengths") else "N/A"
            
            # Gaps
            r_cells[4].text = "• " + "\n• ".join(cand.get("key_gaps", [])) if cand.get("key_gaps") else "N/A"
            
            # Recommendation
            r_cells[5].text = cand.get("executive_summary", "")
            
            for c_i, cell in enumerate(r_cells):
                set_cell_background(cell, "EBF1E5")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = docx.shared.Pt(8.5)

    # 2. Longlist Section (Backup Pool)
    if longlisted:
        h2_2 = doc.add_paragraph()
        h2_2.paragraph_format.space_before = docx.shared.Pt(8)
        h2_2.paragraph_format.space_after = docx.shared.Pt(4)
        run_2 = h2_2.add_run("🟡 Longlist (Backup Pool)")
        run_2.font.name = "Georgia"
        run_2.font.size = docx.shared.Pt(11)
        run_2.font.bold = True
        run_2.font.color.rgb = docx.shared.RGBColor(0, 32, 91)
        
        t_long = doc.add_table(rows=len(longlisted)+1, cols=4)
        t_long.style = 'Table Grid'
        
        headers_long = ["Rank", "Candidate Name", "Score", "Brief Summary / Note"]
        hdr_cells_l = t_long.rows[0].cells
        for i, h in enumerate(headers_long):
            hdr_cells_l[i].text = h
            set_cell_background(hdr_cells_l[i], "C49A45") # Harrow Warm Gold
            p = hdr_cells_l[i].paragraphs[0]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = docx.shared.Pt(8.5)
                r.font.bold = True
                r.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
                
        for l_idx, cand in enumerate(longlisted, 1):
            row_idx = l_idx
            r_cells = t_long.rows[row_idx].cells
            
            # Rank
            r_cells[0].text = str(candidates.index(cand) + 1)
            r_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Name
            p_name = r_cells[1].paragraphs[0]
            add_hyperlink(p_name, cand["rel_path"], cand["candidate_name"])
            
            # Score
            r_cells[2].text = str(cand["score"])
            r_cells[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            r_cells[2].paragraphs[0].runs[0].font.bold = True
            
            # Summary Note
            r_cells[3].text = cand.get("executive_summary", "")
            
            for c_i, cell in enumerate(r_cells):
                set_cell_background(cell, "FEF9E7")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = docx.shared.Pt(8.5)

    # 3. Filtered Out Section (Score < 50)
    if rejected:
        h2_3 = doc.add_paragraph()
        h2_3.paragraph_format.space_before = docx.shared.Pt(8)
        h2_3.paragraph_format.space_after = docx.shared.Pt(2)
        run_3 = h2_3.add_run("🔴 Filtered Out (Score < 50)")
        run_3.font.name = "Georgia"
        run_3.font.size = docx.shared.Pt(11)
        run_3.font.bold = True
        run_3.font.color.rgb = docx.shared.RGBColor(0, 32, 91)
        
        rej_p = doc.add_paragraph()
        rej_p.paragraph_format.space_after = docx.shared.Pt(8)
        bullet_run = rej_p.add_run("• ")
        bullet_run.font.name = "Arial"
        bullet_run.font.size = docx.shared.Pt(8.5)
        
        for idx_r, cand in enumerate(rejected):
            add_hyperlink(rej_p, cand["rel_path"], cand["candidate_name"])
            sc_run = rej_p.add_run(f" ({cand['score']})")
            sc_run.font.name = "Arial"
            sc_run.font.size = docx.shared.Pt(8.5)
            
            if idx_r < len(rejected) - 1:
                sep_run = rej_p.add_run(" | ")
                sep_run.font.name = "Arial"
                sep_run.font.size = docx.shared.Pt(8.5)
                
        reason_run = rej_p.add_run(" — (Unmatched core qualifications)")
        reason_run.font.name = "Arial"
        reason_run.font.size = docx.shared.Pt(8.5)
        reason_run.font.italic = True
        reason_run.font.color.rgb = docx.shared.RGBColor(146, 43, 33)

    # Body Safeguarding Notice
    safeguarding_p = doc.add_paragraph()
    safeguarding_p.paragraph_format.space_before = docx.shared.Pt(12)
    safeguarding_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    safeguarding_run = safeguarding_p.add_run("Harrow International School Bangkok is committed to the safety and protection of children.")
    safeguarding_run.font.name = "Georgia"
    safeguarding_run.font.size = docx.shared.Pt(8.5)
    safeguarding_run.font.italic = True
    safeguarding_run.font.color.rgb = docx.shared.RGBColor(0, 32, 91)

    doc.save(docx_path)
    return docx_path


# ---------------------------------------------------------------------------
# Batch File Template Generator (Dual-Mode EXE/Python)
# ---------------------------------------------------------------------------
def ensure_run_screening_bat(target_dir: Path):
    """Ensure run_screening.bat is present inside target job folder with dual EXE/Python support."""
    bat_path = target_dir / "run_screening.bat"
    bat_content = (
        "@echo off\n"
        "chcp 65001 > nul\n"
        "echo Starting Candidate Screening Engine...\n\n"
        'if exist "%~dp0..\\screener_core.exe" (\n'
        '    "%~dp0..\\screener_core.exe" "%~dp0"\n'
        ') else if exist "%~dp0..\\screener_core.py" (\n'
        '    python "%~dp0..\\screener_core.py" "%~dp0"\n'
        ") else (\n"
        "    echo Error: Neither screener_core.exe nor screener_core.py was found in parent directory.\n"
        ")\n"
        "pause\n"
    )
    bat_path.write_text(bat_content, encoding="utf-8")
    print(f"  [OK] Batch file template created: {bat_path.name}")


# ---------------------------------------------------------------------------
# Main Orchestrator Execution
# ---------------------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print("Usage: python screener_core.py <path_to_job_folder>")
        sys.exit(1)
        
    target_dir_arg = sys.argv[1].strip('"\'').rstrip('\\')
    target_dir = Path(target_dir_arg).resolve()
    
    if not target_dir.exists() or not target_dir.is_dir():
        print(f"Error: Target job folder '{target_dir}' does not exist or is not a directory.")
        sys.exit(1)
        
    print(f"============================================================")
    print(f"🚀 Candidate Screening Engine")
    print(f"Target Job Folder: {target_dir}")
    print(f"============================================================")
    
    # 0. API Key Resolution
    api_key = resolve_api_key(target_dir)
    if not api_key:
        print("❌ Error: GEMINI_API_KEY / GOOGLE_API_KEY environment variable is missing.")
        sys.exit(1)
        
    # 1. Read Job Description
    jd_path, jd_info = find_and_read_job_description(target_dir)
    jd_text = jd_info.get("text", "")
    if not jd_path or (not jd_text.strip() and not jd_info.get("raw_bytes")):
        print(f"❌ Error: No valid Job Description (PDF/DOCX/TXT) found in target folder: {target_dir}")
        sys.exit(1)
    print(f"✔ Job Description Loaded: '{jd_path.relative_to(target_dir)}' ({len(jd_text)} characters)")
    
    # 2. Read Candidate CVs
    cv_items = find_and_read_candidate_cvs(target_dir)
    if not cv_items:
        print(f"❌ Error: No candidate CVs found in subdirectories (e.g. cv/) of: {target_dir}")
        sys.exit(1)
    print(f"✔ Found {len(cv_items)} Candidate CV(s) in subdirectories.")
    
    # 3. Initialize Gemini GenAI Client
    print(f"\n🧠 Initializing Gemini GenAI Client (gemini-2.5-flash)...")
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        print(f"❌ Error initializing GenAI Client: {e}")
        sys.exit(1)
        
    # 4. Evaluate CVs
    print(f"\n🔍 Evaluating candidates against Job Description...")
    evaluations = []
    for idx, cv in enumerate(cv_items, 1):
        print(f"  [{idx}/{len(cv_items)}] Evaluating candidate CV: {cv['rel_path']}...")
        try:
            eval_data = evaluate_cv_against_jd(client, jd_text, cv, jd_info)
            evaluations.append(eval_data)
            print(f"      -> {eval_data['candidate_name']}: Score = {eval_data['score']}")
        except Exception as e:
            print(f"      ❌ Evaluation failed for {cv['rel_path']}: {e}")
            
    if not evaluations:
        print("❌ Error: No candidate evaluations succeeded.")
        sys.exit(1)
        
    # 5. Apply HR Natural Break / Gap Analysis Tiering
    print(f"\n📈 Applying HR Natural Break / Gap Analysis Tiering...")
    tiered_candidates = apply_hr_natural_break_tiering(evaluations)
    
    # Print console summary table
    print("\n------------------------------------------------------------")
    print(f"{'Rank':<5} | {'Candidate Name':<25} | {'Score':<6} | {'Tier':<10}")
    print("------------------------------------------------------------")
    for idx, cand in enumerate(tiered_candidates, 1):
        print(f"{idx:<5} | {cand['candidate_name']:<25} | {cand['score']:<6} | {cand['tier']:<10}")
    print("------------------------------------------------------------\n")
    
    # 6. Export Reports
    print(f"📄 Exporting Candidate Summary Reports...")
    md_file = export_candidate_summary_md(target_dir, jd_path.name, tiered_candidates)
    print(f"  [OK] Exported Markdown Summary: {md_file.name}")
    
    excel_file = export_candidate_summary_excel(target_dir, jd_path.name, tiered_candidates)
    print(f"  [OK] Exported Excel Summary: {excel_file.name}")
    
    docx_file = export_candidate_summary_docx(target_dir, jd_path.name, tiered_candidates)
    print(f"  [OK] Exported Word DOCX Summary: {docx_file.name}")
    
    # 7. Ensure run_screening.bat exists inside target job folder
    ensure_run_screening_bat(target_dir)
    
    print(f"\n✅ Candidate screening completed successfully!")


if __name__ == "__main__":
    main()
